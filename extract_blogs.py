import xml.etree.ElementTree as ET  # Import XML handling library
from docx import Document  # Import library to create Word documents
from bs4 import BeautifulSoup  # Import BeautifulSoup for parsing HTML
import requests  # Import requests to make HTTP requests
from docx.shared import Inches  # Import Inches for specifying image size in the Word document
from io import BytesIO  # Import BytesIO for handling image data in memory
import html  # Import html to unescape HTML entities

# Load XML file
xml_file = "travel-blog-11-09-2024.xml"  # Path to the XML file containing blog entries
tree = ET.parse(xml_file)  # Parse the XML file
tree.getroot()  # Get the root element of the XML
root = tree.getroot()  # Get the root element of the XML

# Namespaces used in the XML
namespaces = {
    'atom': 'http://www.w3.org/2005/Atom',  # Atom namespace for XML elements
    'georss': 'http://www.georss.org/georss',  # GeoRSS namespace for geolocation elements
    'gd': 'http://schemas.google.com/g/2005',  # Google namespace for additional elements
    'thr': 'http://purl.org/syndication/thread/1.0'  # Thread namespace for comments and thread information
}

# Create a new Word document
doc = Document()  # Create a new Word document
doc.add_heading("First safari: blogs", level=1)  # Add a main heading to the document

first_time = True # for the first level 2 header, we dont need a page break

# Iterate through each <entry> element and extract relevant information
for entry in root.findall('atom:entry', namespaces):  # Find all <entry> elements using the Atom namespace
    entry_id = entry.find('atom:id', namespaces).text if entry.find('atom:id', namespaces) is not None else ''  # Get the <id> element text
    if entry_id.startswith("tag:blogger.com,1999:blog-7121986992391433647.post"):  # Filter entries based on a specific blog ID prefix
        title = entry.find('atom:title', namespaces).text if entry.find('atom:title', namespaces) is not None else 'No Title'  # Get the <title> element text
        if title.startswith("Botswana 2024"):  # Further filter entries based on title
            published = entry.find('atom:published', namespaces).text if entry.find('atom:published', namespaces) is not None else 'No Date'  # Get the <published> element text
            content = entry.find('atom:content', namespaces).text if entry.find('atom:content', namespaces) is not None else 'No Content'  # Get the <content> element text
            author = entry.find('atom:author/atom:name', namespaces).text if entry.find('atom:author/atom:name', namespaces) is not None else 'No Author'  # Get the author name
            
            # Replace HTML entities (&lt;, &gt;, etc.) with their corresponding characters
            # content = html.unescape(content)  # Unescape HTML entities in the content
            
            # Use BeautifulSoup to parse content while retaining structure
            soup = BeautifulSoup(content, 'html.parser')  # Parse the content using BeautifulSoup
            
            # Add entry to Word document
            if first_time:
                first_time = False
            else:
                doc.add_page_break()  # Start a new page for each blog entry
            doc.add_heading(title, level=2)  # Add the title of the blog entry as a heading
            doc.add_paragraph(f'Published: {published}')  # Add the publication date
            doc.add_paragraph(f'Author: {author}')  # Add the author name
            
            # Iterate through the parsed HTML to retain headings, paragraphs, images, etc.
            for element in soup.descendants:  # Iterate through all descendants of the parsed content
                if element.name == 'h3':  # If the element is an <h3> heading
                    doc.add_heading(element.get_text(), level=3)  # Add the <h3> text as a level-3 heading
                elif element.name == 'p':  # If the element is a <p> paragraph
                    # Add paragraph text if present
                    if element.get_text(strip=True):  # Check if the paragraph contains text
                        doc.add_paragraph(element.get_text())  # Add the paragraph text to the document
                elif element.name == 'img':  # If the element is an <img> tag
                    # Check for image links
                    img_src = element.get('src')  # Get the 'src' attribute of the image
                    if img_src and img_src.startswith("https://blogger.googleusercontent.com/"):  # Only process specific image URLs
                        try:
                            response = requests.get(img_src)  # Make an HTTP request to get the image
                            if response.status_code == 200:  # Check if the request was successful
                                image_stream = BytesIO(response.content)  # Create a BytesIO object from the image content
                                doc.add_picture(image_stream, width=Inches(4.0))  # Add the image to the document with specified width
                                
                                # Add the image alt and title text if available
                                alt_text = element.get('alt', 'No description available')  # Get the 'alt' attribute of the image
                                title_text = element.get('title', 'No title available')  # Get the 'title' attribute of the image
                                doc.add_paragraph(f'Image description: {alt_text}')  # Add the alt text below the image
                                doc.add_paragraph(f'Image title: {title_text}')  # Add the title text below the image
                        except Exception as e:  # Handle any errors that occur during the request
                            print(f"Failed to download image: {img_src}, error: {e}")  # Print an error message

# Save the Word document
output_file = 'blog_entries_v04.docx' # Path to save the Word document
doc.save(output_file)  # Save the Word document

output_file  # Return the path of the saved Word document
