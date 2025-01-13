import json
from googleapiclient.discovery import build
from docx import Document
import requests
from bs4 import BeautifulSoup
from typing import List
from io import BytesIO
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import os
import time
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class TravelBlogExtractor:
    def __init__(self, 
                 blog_post_list=".\\test_output\\blog_post_urls.txt",
                 page_load_wait=20, 
                 output_docx_file=".\\test_output\\travel_blog_posts.docx", 
                 output_pdf_file=".\\test_output\\travel_blog_posts.pdf", 
                 output_docx_path=".\\test_output", 
                 file_name_starts_with="travel_blog_posts_",
                 chunk_size=1):
        self.config_file = ".\\config.json"
        self.blog_post_list = blog_post_list
        self.page_load_wait = page_load_wait
        self.output_docx_file = output_docx_file
        self.output_pdf_file = output_pdf_file
        self.output_docx_path = output_docx_path
        self.file_name_starts_with = file_name_starts_with
        self.chunk_size = chunk_size

        with open(self.config_file, "r") as config_file:
            self.config = json.load(config_file)
        self.blogger_api_key = self.config['BLOGGER_API_KEY']
        self.travel_blog_id = self.config['TRAVEL_BLOG_ID']
        self.gmaps_api_key = self.config['GMAPS_API_KEY']

    def get_travel_blog_urls(self) -> List[str]:
        service = build('blogger', 'v3', developerKey=self.blogger_api_key)
        post_links = []
        request = service.posts().list(blogId=self.travel_blog_id, maxResults=100)
        while request is not None:
            response = request.execute()
            for post in response.get('items', []):
                post_url = post.get('url', '').lower()
                if 'second-post' not in post_url and 'first-post' not in post_url:
                    post_links.append(post['url'])
            request = service.posts().list_next(request, response)
        return list(reversed(post_links))

    def add_caption(self, doc, text):
        """
        Adds a caption paragraph to the document.

        :param doc: The Word document object.
        :param text: The caption text to add.
        """
        caption_paragraph = doc.add_paragraph(text, style='Caption')
        caption_format = caption_paragraph.paragraph_format
        caption_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the caption
        caption_format.space_before = Pt(0)  # Optional: Space before the caption
        caption_format.space_after = Pt(6)   # Optional: Space after the caption

    def add_centered_image(self, doc, img_src):
        """
        Adds a centered image to the document.

        :param doc: The Word document object.
        :param img_src: The source URL or path of the image to add.
        """
        try:
            img_response = requests.get(img_src)
            if img_response.status_code == 200:
                image_stream = BytesIO(img_response.content)
                paragraph = self.add_formatted_paragraph(doc,
                                                            "",
                                                            style='Body Text',
                                                            before=4,
                                                            after=4,
                                                            space_between=False)
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6.0))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                print(f"Failed to download image: {img_src} - Status Code: {img_response.status_code}")
        except Exception as img_e:
            print(f"Failed to add image: {img_src}, error: {img_e}")

    def download_and_add_image(self, doc, img_src, element):
        """
        Downloads an image from the provided source URL and adds it to the document,
        with an optional caption derived from the element attributes.
        """
        try:
            self.add_centered_image(doc, img_src)

            # Retrieve caption from the image element
            caption = element.get('title') or element.get('alt', '').strip()
            if caption:
                self.add_caption(doc, caption)
        except Exception as img_e:
            print(f"Failed to download or add image: {img_src}, error: {img_e}")

    def download_and_add_map_sshot(self, doc, title, map_src):
        try:
            options = webdriver.ChromeOptions()
            options.add_argument('--headless')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

            print(f"Loading map at {map_src} in page {title}")
            driver.get(map_src)

            # Wait for the page to load completely
            time.sleep(self.page_load_wait)

            # Wait until the map iframe is loaded
            # WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.TAG_NAME, 'iframe')))

            screenshot = driver.get_screenshot_as_png()
            map_stream = BytesIO(screenshot)
            doc.add_picture(map_stream, width=Inches(6.0))
            driver.quit()
        except Exception as map_e:
            print(f"Failed to download map: {map_src}, error: {map_e}")

    @staticmethod
    def set_font_to_aptos(doc):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Aptos'
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Aptos')
                rFonts.set(qn('w:hAnsi'), 'Aptos')
                rFonts.set(qn('w:eastAsia'), 'Aptos')
                rFonts.set(qn('w:cs'), 'Aptos')
                rPr.append(rFonts)

    def add_formatted_paragraph(self, doc, text, style=None, before=3, after=3, space_between=False):
        """
        Adds a formatted paragraph to the document.

        :param doc: The Word document object.
        :param text: The text to add to the paragraph.
        :param style: Optional style to apply to the paragraph.
        :param before: Space before the paragraph (in points).
        :param after: Space after the paragraph (in points).
        :param space_between: Don't add space between paragraphs of the same style (True/False)
        """
        paragraph = doc.add_paragraph(text, style=style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before) # 3pt before the paragraph
        paragraph_format.space_after = Pt(after) # 3pt after the paragraph
        paragraph_format.space_between = space_between  # Uncheck "Don't add space between paragraphs of the same style"
        return paragraph

    def process_list(self, doc, list_element, level=0):
        """
        Processes <ul> or <ol> elements and their <li> children hierarchically,
        ensuring proper handling of nested lists with appropriate indentation.
        """
        # Map levels to corresponding Word styles for unordered and ordered lists
        ul_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3', 'List Bullet 4', 'List Bullet 5']
        ol_styles = ['List Number', 'List Number 2', 'List Number 3', 'List Number 4', 'List Number 5']

        # Determine the style set based on the type of list
        if list_element.name == 'ul':
            list_styles = ul_styles
        elif list_element.name == 'ol':
            list_styles = ol_styles
        else:
            return  # Ignore unsupported elements

        current_style = list_styles[min(level, len(list_styles) - 1)]  # Use the deepest style if level exceeds limit

        # Loop through each child of the list element
        for child in list_element.children:
            if child.name == 'li':  # Process list items
                li_text = ""  # Initialize text for the list item

                # Process the contents of the <li> (text, links, or spans)
                for li_child in child.children:
                    if li_child.name == 'b':  # Handle bold text explicitly
                        bold_text = li_child.get_text(strip=True)
                        li_text += f"{bold_text}"  # Add bold text (adjust formatting as needed)
                    elif li_child.name == 'a':  # Handle hyperlinks
                        href = li_child.get('href', '')
                        link_text = li_child.get_text(strip=True)
                        li_text += f"{link_text} ({href}) "
                    elif li_child.string:  # Handle plain text
                        li_text += li_child.string

                # Add the formatted paragraph for the <li>
                self.add_formatted_paragraph(doc, li_text.strip(), style=current_style, before=3, after=3)

                # Handle nested <ul> or <ol>
                for li_child in child.children:
                    if li_child.name in ['ul', 'ol']:
                        self.process_list(doc, li_child, level=level + 1)

            elif child.name in ['ul', 'ol']:  # Nested <ul> or <ol> directly under a list (rare case)
                self.process_list(doc, child, level=level + 1)

    def process_blog_post(self, doc, link):
        print("++ entering process blog post ++")
        try:
            response = requests.get(link)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                title = soup.find("title").get_text().replace("Travel diaries: ", "") if soup.find("title") else "No Title"
                doc.add_heading(title, level=2)

                post_body = soup.find("div", class_="post-body")
                # print(post_body)
                if post_body:
                    for element in post_body.descendants:

                        if element.name == 'h1':
                            doc.add_heading(element.get_text(), level=1)
                        elif element.name == 'h2':
                            doc.add_heading(element.get_text(), level=2)
                        elif element.name == 'h3':
                            doc.add_heading(element.get_text(), level=3)
                        elif element.name == 'h4':
                            doc.add_heading(element.get_text(), level=4)

                        elif element.name == 'p':
                            # if element.get_text(strip=True):
                            #     print("element.get_text() = ", element.get_text())
                            #     doc.add_paragraph(element.get_text())
                            paragraph = self.add_formatted_paragraph(doc,
                                                                     "",
                                                                     style='Body Text',
                                                                     before=4,
                                                                     after=4,
                                                                     space_between=False)
                            # paragraph = doc.add_paragraph()  # Create a paragraph in the docx
                            for child in element.children:  # Traverse the direct children of <p>
                                if child.name == 'a':  # Handle hyperlinks
                                    href = child.get('href', '')
                                    link_text = child.get_text(strip=True)
                                    run = paragraph.add_run(link_text)
                                    # run.font.underline = True  # Make the text underlined
                                    # run.font.color.rgb = RGBColor(0, 0, 255)  # Make the text blue
                                    paragraph.add_run(f" ({href})")  # Append the URL
                                elif child.name == 'span':  # Handle spans and their content
                                    for span_child in child.children:
                                        if span_child.name == 'a':  # Handle links inside spans
                                            href = span_child.get('href', '')
                                            link_text = span_child.get_text(strip=True)
                                            run = paragraph.add_run(link_text)
                                            # run.font.underline = True
                                            # run.font.color.rgb = RGBColor(0, 0, 255)
                                            paragraph.add_run(f" ({href})")
                                        elif span_child.string:  # Handle plain text inside spans
                                            paragraph.add_run(span_child.string)
                                elif child.string:  # Add plain text
                                    paragraph.add_run(child.string)

                        elif element.name == 'ul':  # Process unordered lists
                            self.process_list(doc, element)

                        elif element.name == 'ol':
                            self.process_list(doc, element)

                        elif element.name == 'img':
                            img_src = element.get('src')
                            if img_src:
                                self.download_and_add_image(doc, img_src, element)
                        elif element.name == 'iframe':
                            map_src = element.get('src')
                            if map_src and map_src.startswith("https://www.google.com/maps"):
                                self.download_and_add_map_sshot(doc, title, map_src)
            else:
                print(f"Failed to load post: {link} - Status Code: {response.status_code}")
        except Exception as e:
            print(f"Error processing {link}: {e}")

    def create_travel_blog_docx(self, output_docx_path, blog_post_list):
        with open(blog_post_list, "r", encoding="utf-8") as file:
            post_links = [line.strip() for line in file.readlines()]

        doc = Document()
        doc.add_heading("Travel Blog Posts", level=1)

        for idx, link in enumerate(post_links):
            if idx > 0:
                doc.add_page_break()
            self.process_blog_post(doc, link)

        self.set_font_to_aptos(doc)
        doc.save(output_docx_path)

    def create_travel_blog_docx_split(self, output_docx_path, blog_post_list):
        with open(blog_post_list, "r", encoding="utf-8") as file:
            post_links = [line.strip() for line in file.readlines()]

        chunks = [post_links[i:i + self.chunk_size] for i in range(0, len(post_links), self.chunk_size)]

        for idx, chunk in enumerate(chunks):
            doc = Document()

            for link in chunk:
                self.process_blog_post(doc, link)

            self.set_font_to_aptos(doc)
            doc_name = os.path.join(output_docx_path, f'travel_blog_posts_{idx + 1:02}.docx')
            doc.save(doc_name)

    @staticmethod
    def convert_docx_to_pdf(docx_file_path, pdf_file_path) -> str:
        try:
            convert(docx_file_path, pdf_file_path)
            return f"Successfully converted {docx_file_path} to {pdf_file_path}"
        except Exception as e:
            return f"Failed to convert {docx_file_path} to PDF: {e}"

    @staticmethod
    def convert_docx_to_pdf_multi(docx_path, pdf_path, file_name_starts_with):
        try:
            for file_name in os.listdir(docx_path):
                if file_name.startswith(file_name_starts_with) and file_name.endswith('.docx'):
                    docx_file = os.path.join(docx_path, file_name)
                    pdf_file = os.path.join(pdf_path, f"{os.path.splitext(file_name)[0]}.pdf")
                    convert(docx_file, pdf_file)
                    print(f"Successfully converted {docx_file} to {pdf_file}")
        except Exception as e:
            print(f"Failed to convert documents starting with {file_name_starts_with} to PDFs: {e}")

def log_execution(task_name, func, *args, **kwargs):
    start_time = datetime.now()
    print(f"{start_time}: Starting {task_name}")
    result = func(*args, **kwargs)
    end_time = datetime.now()
    print(f"{end_time}: Completed {task_name} (Time taken: {(end_time - start_time).total_seconds()} seconds)")
    return result

if __name__ == "__main__":
    # doc = Document()
    # print([style.name for style in doc.styles])

    extractor = TravelBlogExtractor()

    post_links = log_execution("get_travel_blog_urls", extractor.get_travel_blog_urls)

    log_execution(
        "Writing blog post URLs to file",
        lambda: open(extractor.blog_post_list, "w", encoding="utf-8").writelines([link + "\n" for link in post_links])
    )

    log_execution("create_travel_blog_docx", extractor.create_travel_blog_docx, extractor.output_docx_file, extractor.blog_post_list)

    log_execution("convert_docx_to_pdf", extractor.convert_docx_to_pdf, extractor.output_docx_file, extractor.output_pdf_file)

    log_execution("create_travel_blog_docx_split", extractor.create_travel_blog_docx_split, extractor.output_docx_path, extractor.blog_post_list)

    log_execution("convert_docx_to_pdf_multi", extractor.convert_docx_to_pdf_multi, extractor.output_docx_path, extractor.output_docx_path, extractor.file_name_starts_with)
